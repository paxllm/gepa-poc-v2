"""
Multi-format resume text extraction.
Supports PDF (via PyMuPDF), DOCX (via python-docx), and TXT files.
"""

import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


def parse_resume(file_path: str | Path) -> str:
    """
    Extract text from a resume file.

    Args:
        file_path: Path to the resume file (PDF, DOCX, or TXT).

    Returns:
        Cleaned, normalized text content.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif suffix == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: .pdf, .docx, .txt")


def _parse_pdf(file_path: Path) -> str:
    """Extract text from PDF using PyMuPDF."""
    text_parts = []
    with fitz.open(str(file_path)) as doc:
        for page in doc:
            text_parts.append(page.get_text())

    raw_text = "\n".join(text_parts)
    return _normalize_text(raw_text)


def _parse_docx(file_path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    doc = Document(str(file_path))
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    raw_text = "\n".join(text_parts)
    return _normalize_text(raw_text)


def _parse_txt(file_path: Path) -> str:
    """Read plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        raw_text = f.read()
    return _normalize_text(raw_text)


def _normalize_text(text: str) -> str:
    """
    Normalize extracted text:
    - Collapse multiple whitespace
    - Remove excessive newlines
    - Strip leading/trailing whitespace
    """
    # Replace tabs with spaces
    text = text.replace("\t", " ")

    # Collapse multiple spaces into one
    text = re.sub(r" {2,}", " ", text)

    # Collapse 3+ newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip each line
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    return text.strip()
